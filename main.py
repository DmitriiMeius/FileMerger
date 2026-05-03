# FileMerger PRO v3.6 OCR Full Language Selector
# Copyright (c) 2026 Universal Experts, LLC

from pathlib import Path
from datetime import datetime
from contextlib import redirect_stderr

import os
import re
import csv
import sys
import zipfile
import threading
import queue
import unicodedata
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import fitz
import pandas as pd
from docx import Document
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text
from charset_normalizer import from_bytes

import pytesseract
from PIL import Image


APP_NAME = "FileMerger PRO"
APP_VERSION = "v3.6 OCR Full Languages"
APP_COPYRIGHT = "© 2026 Universal Experts, LLC"

APP_DIR = Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parent
BUNDLE_DIR = Path(getattr(sys, "_MEIPASS", APP_DIR))


def resource_path(*parts):
    app_path = APP_DIR.joinpath(*parts)
    if app_path.exists():
        return app_path

    return BUNDLE_DIR.joinpath(*parts)


LOCAL_TESSERACT = resource_path("Tesseract-OCR", "tesseract.exe")
if LOCAL_TESSERACT.exists():
    pytesseract.pytesseract.tesseract_cmd = str(LOCAL_TESSERACT)

ICON_FILE = resource_path("FileMerger_PRO.ico")

STANDALONE_FORMATS = [
    "*.txt", "*.rtf", "*.pdf", "*.docx",
    "*.xls", "*.xlsx", "*.fb2", "*.zip",
    "*.epub", "*.html", "*.htm", "*.md", "*.csv"
]

OCR_LANG_OPTIONS_MAIN = {
    "English": "eng",
    "Russian": "rus",
    "Spanish": "spa",
}

OCR_LANG_NAMES = {
    "afr": "Afrikaans",
    "amh": "Amharic",
    "ara": "Arabic",
    "asm": "Assamese",
    "aze": "Azerbaijani",
    "aze_cyrl": "Azerbaijani Cyrillic",
    "bel": "Belarusian",
    "ben": "Bengali",
    "bod": "Tibetan",
    "bos": "Bosnian",
    "bre": "Breton",
    "bul": "Bulgarian",
    "cat": "Catalan",
    "ceb": "Cebuano",
    "ces": "Czech",
    "chi_sim": "Chinese Simplified",
    "chi_sim_vert": "Chinese Simplified Vertical",
    "chi_tra": "Chinese Traditional",
    "chi_tra_vert": "Chinese Traditional Vertical",
    "chr": "Cherokee",
    "cos": "Corsican",
    "cym": "Welsh",
    "dan": "Danish",
    "deu": "German",
    "deu_latf": "German Fraktur",
    "div": "Divehi",
    "dzo": "Dzongkha",
    "ell": "Greek",
    "eng": "English",
    "enm": "Middle English",
    "epo": "Esperanto",
    "equ": "Math / Equation",
    "est": "Estonian",
    "eus": "Basque",
    "fao": "Faroese",
    "fas": "Persian",
    "fil": "Filipino",
    "fin": "Finnish",
    "fra": "French",
    "frm": "Middle French",
    "fry": "Western Frisian",
    "gla": "Scottish Gaelic",
    "gle": "Irish",
    "glg": "Galician",
    "grc": "Ancient Greek",
    "guj": "Gujarati",
    "hat": "Haitian Creole",
    "heb": "Hebrew",
    "hin": "Hindi",
    "hrv": "Croatian",
    "hun": "Hungarian",
    "hye": "Armenian",
    "iku": "Inuktitut",
    "ind": "Indonesian",
    "isl": "Icelandic",
    "ita": "Italian",
    "ita_old": "Old Italian",
    "jav": "Javanese",
    "jpn": "Japanese",
    "jpn_vert": "Japanese Vertical",
    "kan": "Kannada",
    "kat": "Georgian",
    "kat_old": "Old Georgian",
    "kaz": "Kazakh",
    "khm": "Khmer",
    "kir": "Kyrgyz",
    "kmr": "Northern Kurdish",
    "kor": "Korean",
    "lao": "Lao",
    "lat": "Latin",
    "lav": "Latvian",
    "lit": "Lithuanian",
    "ltz": "Luxembourgish",
    "mal": "Malayalam",
    "mar": "Marathi",
    "mkd": "Macedonian",
    "mlt": "Maltese",
    "mon": "Mongolian",
    "mri": "Maori",
    "msa": "Malay",
    "mya": "Burmese",
    "nep": "Nepali",
    "nld": "Dutch",
    "nor": "Norwegian",
    "oci": "Occitan",
    "ori": "Odia",
    "pan": "Punjabi",
    "pol": "Polish",
    "por": "Portuguese",
    "pus": "Pashto",
    "que": "Quechua",
    "ron": "Romanian",
    "rus": "Russian",
    "san": "Sanskrit",
    "sin": "Sinhala",
    "slk": "Slovak",
    "slv": "Slovenian",
    "snd": "Sindhi",
    "spa": "Spanish",
    "spa_old": "Old Spanish",
    "sqi": "Albanian",
    "srp": "Serbian",
    "srp_latn": "Serbian Latin",
    "sun": "Sundanese",
    "swa": "Swahili",
    "swe": "Swedish",
    "syr": "Syriac",
    "tam": "Tamil",
    "tat": "Tatar",
    "tel": "Telugu",
    "tgk": "Tajik",
    "tha": "Thai",
    "tir": "Tigrinya",
    "ton": "Tongan",
    "tur": "Turkish",
    "uig": "Uyghur",
    "ukr": "Ukrainian",
    "urd": "Urdu",
    "uzb": "Uzbek",
    "uzb_cyrl": "Uzbek Cyrillic",
    "vie": "Vietnamese",
    "yid": "Yiddish",
    "yor": "Yoruba",
}


LANG = {
    "en": {
        "subtitle": "Standalone Document Cleaner, Converter & OCR",
        "formats": "Supported: TXT, RTF, PDF, DOCX, XLS, XLSX,\nFB2, FB2.ZIP, EPUB, HTML, HTM, MD, CSV",
        "ocr_title": "OCR languages",
        "more_langs": "MORE OCR LANGUAGES",
        "hide_langs": "HIDE OCR LANGUAGES",
        "select": "SELECT FILES",
        "convert": "CONVERT TO TXT",
        "ready": "Ready.",
        "new_session": "New session ready.",
        "selected": "Selected files",
        "no_files_title": "No files",
        "no_files_msg": "Select files first.",
        "no_ocr_title": "No OCR language",
        "no_ocr_msg": "Select at least one OCR language.",
        "save_title": "Save output file",
        "text_files": "Text files",
        "supported_files": "Supported files",
        "all_files": "All files",
        "working": "Working",
        "idle": "Idle",
        "completed": "Completed",
        "done": "Done",
        "file_converted": "File converted",
        "open_folder": "Open folder?",
        "next": "Next",
        "convert_more": "Convert more files?",
        "progress": "Progress",
        "remaining": "Remaining",
        "adding": "Adding",
        "empty": "EMPTY",
        "skipped": "SKIPPED",
        "error": "Error",
    },
    "ru": {
        "subtitle": "Автономный очиститель, конвертер и OCR",
        "formats": "Поддерживается: TXT, RTF, PDF, DOCX, XLS, XLSX,\nFB2, FB2.ZIP, EPUB, HTML, HTM, MD, CSV",
        "ocr_title": "Языки OCR",
        "more_langs": "ЕЩЁ ЯЗЫКИ OCR",
        "hide_langs": "СКРЫТЬ ЯЗЫКИ OCR",
        "select": "ВЫБРАТЬ ФАЙЛЫ",
        "convert": "КОНВЕРТИРОВАТЬ В TXT",
        "ready": "Готово.",
        "new_session": "Новая сессия готова.",
        "selected": "Выбрано файлов",
        "no_files_title": "Нет файлов",
        "no_files_msg": "Сначала выбери файлы.",
        "no_ocr_title": "Нет языка OCR",
        "no_ocr_msg": "Выбери хотя бы один язык OCR.",
        "save_title": "Сохранить готовый файл",
        "text_files": "Текстовые файлы",
        "supported_files": "Поддерживаемые файлы",
        "all_files": "Все файлы",
        "working": "Работает",
        "idle": "Ожидание",
        "completed": "Завершено",
        "done": "Готово",
        "file_converted": "Файл сконвертирован",
        "open_folder": "Перейти в папку с файлом?",
        "next": "Дальше",
        "convert_more": "Конвертировать ещё файлы?",
        "progress": "Прогресс",
        "remaining": "Осталось",
        "adding": "Добавляю",
        "empty": "ПУСТО",
        "skipped": "ПРОПУЩЕН",
        "error": "Ошибка",
    },
    "es": {
        "subtitle": "Limpiador, convertidor y OCR autónomo",
        "formats": "Soporta: TXT, RTF, PDF, DOCX, XLS, XLSX,\nFB2, FB2.ZIP, EPUB, HTML, HTM, MD, CSV",
        "ocr_title": "Idiomas OCR",
        "more_langs": "MÁS IDIOMAS OCR",
        "hide_langs": "OCULTAR IDIOMAS OCR",
        "select": "SELECCIONAR ARCHIVOS",
        "convert": "CONVERTIR A TXT",
        "ready": "Listo.",
        "new_session": "Nueva sesión lista.",
        "selected": "Archivos seleccionados",
        "no_files_title": "Sin archivos",
        "no_files_msg": "Primero selecciona archivos.",
        "no_ocr_title": "Sin idioma OCR",
        "no_ocr_msg": "Selecciona al menos un idioma OCR.",
        "save_title": "Guardar archivo de salida",
        "text_files": "Archivos de texto",
        "supported_files": "Archivos soportados",
        "all_files": "Todos los archivos",
        "working": "Trabajando",
        "idle": "En espera",
        "completed": "Completado",
        "done": "Listo",
        "file_converted": "Archivo convertido",
        "open_folder": "¿Abrir carpeta del archivo?",
        "next": "Siguiente",
        "convert_more": "¿Convertir más archivos?",
        "progress": "Progreso",
        "remaining": "Restantes",
        "adding": "Agregando",
        "empty": "VACÍO",
        "skipped": "OMITIDO",
        "error": "Error",
    }
}


def get_installed_ocr_languages():
    tessdata_dir = resource_path("Tesseract-OCR", "tessdata")

    if not tessdata_dir.exists():
        tessdata_dir = Path(r"C:\Program Files\Tesseract-OCR\tessdata")

    langs = []

    if tessdata_dir.exists():
        for file in tessdata_dir.glob("*.traineddata"):
            code = file.stem

            if code == "osd":
                continue

            name = OCR_LANG_NAMES.get(code, code)
            langs.append((name, code))

    return sorted(langs, key=lambda x: x[0].lower())


def make_output_filename(files=None):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if files and len(files) == 1:
        return f"{files[0].stem}.txt"

    if files and len(files) > 1:
        return f"{files[0].stem}_merged_{timestamp}.txt"

    return f"merged_{timestamp}.txt"


def natural_key(path: Path):
    nums = re.findall(r"\d+", path.name)
    return [int(n) for n in nums] if nums else [999999]


def decode_bytes_smart(data: bytes) -> str:
    if not data:
        return ""

    result = from_bytes(data).best()

    if result:
        try:
            return str(result)
        except Exception:
            pass

    for enc in ["utf-8-sig", "utf-8", "cp1251", "windows-1251", "cp1252", "latin1", "utf-16"]:
        try:
            return data.decode(enc)
        except Exception:
            continue

    return data.decode("utf-8", errors="ignore")


def fix_mojibake(text: str) -> str:
    def fix_latin1_cp1251(match):
        fragment = match.group(0)
        try:
            return fragment.encode("latin1").decode("cp1251")
        except Exception:
            return fragment

    return re.sub(r"[\u00A1-\u00FF]{2,}", fix_latin1_cp1251, text)


def remove_hidden_unicode(text: str) -> str:
    cleaned = []

    for char in text:
        category = unicodedata.category(char)

        if char in ["\n", "\t"]:
            cleaned.append(char)
            continue

        if category in ["Cf", "Cc"]:
            continue

        cleaned.append(char)

    return "".join(cleaned)


def smart_clean_text(text: str) -> str:
    text = fix_mojibake(text)

    weird_spaces = {
        "\u00A0": " ",
        "\u202F": " ",
        "\u2009": " ",
        "\u2002": " ",
        "\u2003": " ",
        "\u2004": " ",
        "\u2005": " ",
        "\u2006": " ",
        "\u2007": " ",
        "\u2008": " ",
        "\u200A": " ",
        "\u00AD": "",
        "\u200B": "",
        "\u200C": "",
        "\u200D": "",
        "\u2060": "",
        "\u034F": "",
        "\u180E": "",
        "\uFEFF": "",
    }

    for bad, good in weird_spaces.items():
        text = text.replace(bad, good)

    text = remove_hidden_unicode(text)

    replacements = {
        "“": '"',
        "”": '"',
        "„": '"',
        "«": '"',
        "»": '"',
        "‘": "'",
        "’": "'",
        "–": "-",
        "—": "-",
        "…": "...",
    }

    for bad, good in replacements.items():
        text = text.replace(bad, good)

    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"(?<![.!?:;])\n(?!\n)", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+([,.!?;:])", r"\1", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def looks_like_page_number(text: str) -> bool:
    return bool(re.fullmatch(r"\d{1,4}", text.strip()))


def looks_like_reference(text: str) -> bool:
    s = text.strip()

    if re.match(r"^(\d{1,3}|[*†‡])[\).]\s+\S+", s):
        return True

    if re.match(r"^\[\d{1,3}\]\s+\S+", s):
        return True

    if re.search(
        r"\b(doi|isbn|journal|publisher|press|vol\.|pp\.|стр\.|с\.|изд\.|том|ред\.)\b",
        s,
        re.IGNORECASE,
    ):
        return True

    return False


def normalize_inline_reference_markers(text: str) -> str:
    return re.sub(
        r"(?<=[A-Za-zА-Яа-яЁё])(\d{1,3})(?=[\.\,\;\:\!\?\s])",
        r"[\1]",
        text
    )


def is_probable_title(text: str, font_size: float, avg_size: float) -> bool:
    s = text.strip()

    if len(s) > 120:
        return False

    if font_size >= avg_size * 1.25:
        return True

    if s.isupper() and len(s) > 4:
        return True

    return False


def read_txt(path: Path):
    return decode_bytes_smart(path.read_bytes())


def read_md(path: Path):
    return decode_bytes_smart(path.read_bytes())


def read_rtf(path: Path):
    content = decode_bytes_smart(path.read_bytes())
    return rtf_to_text(content)


def read_docx(path: Path):
    doc = Document(path)
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_excel(path: Path):
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    parts = []

    for sheet_name, df in sheets.items():
        parts.append(f"\n--- SHEET: {sheet_name} ---\n")
        parts.append(df.fillna("").to_string(index=False))

    return "\n".join(parts)


def read_csv_file(path: Path):
    text = decode_bytes_smart(path.read_bytes())
    rows = []

    try:
        dialect = csv.Sniffer().sniff(text[:4096])
    except Exception:
        dialect = csv.excel

    for row in csv.reader(text.splitlines(), dialect):
        rows.append(" | ".join(row))

    return "\n".join(rows)


def read_html(path: Path):
    content = decode_bytes_smart(path.read_bytes())
    soup = BeautifulSoup(content, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    return soup.get_text("\n", strip=True)


def read_epub(path: Path):
    parts = []

    with zipfile.ZipFile(path, "r") as z:
        html_files = [
            name for name in z.namelist()
            if name.lower().endswith((".xhtml", ".html", ".htm"))
        ]

        html_files = sorted(html_files, key=lambda x: natural_key(Path(x)))

        for name in html_files:
            with z.open(name) as f:
                content = decode_bytes_smart(f.read())
                soup = BeautifulSoup(content, "html.parser")

                for tag in soup(["script", "style", "noscript"]):
                    tag.decompose()

                text = soup.get_text("\n", strip=True)

                if text:
                    parts.append(text)

    return "\n\n".join(parts)


def ocr_pdf_page(page, ocr_langs):
    pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    return pytesseract.image_to_string(
        img,
        lang=ocr_langs,
        config="--oem 3 --psm 6"
    )


def read_pdf_structural(path: Path, ocr_langs: str):
    doc = fitz.open(path)
    main_parts = []
    references = []
    repeated_candidates = {}

    with open(os.devnull, "w") as devnull:
        with redirect_stderr(devnull):
            for page_index, page in enumerate(doc, start=1):
                page_h = page.rect.height
                data = page.get_text("dict")
                page_items = []
                font_sizes = []

                plain_check = page.get_text("text").strip()

                if len(plain_check) < 30:
                    try:
                        ocr_text = ocr_pdf_page(page, ocr_langs)
                        ocr_text = smart_clean_text(ocr_text)

                        if ocr_text:
                            main_parts.append(ocr_text)
                            continue
                    except Exception:
                        continue

                for block in data.get("blocks", []):
                    if "lines" not in block:
                        continue

                    for line in block["lines"]:
                        for span in line.get("spans", []):
                            txt = span.get("text", "").strip()
                            size = span.get("size", 0)

                            if txt:
                                font_sizes.append(size)

                avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 10

                for block in data.get("blocks", []):
                    if "lines" not in block:
                        continue

                    x0, y0, x1, y1 = block["bbox"]

                    block_text_parts = []
                    max_font = 0

                    for line in block["lines"]:
                        line_parts = []

                        for span in line.get("spans", []):
                            txt = span.get("text", "")
                            size = span.get("size", 0)
                            max_font = max(max_font, size)

                            if txt.strip():
                                line_parts.append(txt)

                        if line_parts:
                            block_text_parts.append("".join(line_parts).strip())

                    block_text = "\n".join(block_text_parts).strip()

                    if not block_text:
                        continue

                    top_zone = y0 < page_h * 0.07
                    bottom_zone = y1 > page_h * 0.93
                    footnote_zone = y0 > page_h * 0.74

                    if looks_like_page_number(block_text):
                        continue

                    if top_zone or bottom_zone:
                        key = re.sub(r"\d+", "", block_text).strip().lower()
                        if key:
                            repeated_candidates[key] = repeated_candidates.get(key, 0) + 1
                        continue

                    if footnote_zone and looks_like_reference(block_text):
                        references.append((page_index, block_text))
                        continue

                    block_text = normalize_inline_reference_markers(block_text)

                    if is_probable_title(block_text, max_font, avg_size):
                        page_items.append(f"\n\n## {block_text}\n")
                    else:
                        page_items.append(block_text)

                if page_items:
                    main_parts.append("\n".join(page_items))

    doc.close()

    main_text = "\n\n".join(main_parts)

    for repeated, count in repeated_candidates.items():
        if count >= 3 and len(repeated) > 4:
            pattern = re.escape(repeated)
            main_text = re.sub(pattern, "", main_text, flags=re.IGNORECASE)

    main_text = smart_clean_text(main_text)

    if references:
        refs = ["\n\n===== EXTRACTED REFERENCES / FOOTNOTES =====\n"]

        for i, (page, ref_text) in enumerate(references, start=1):
            refs.append(f"[{i}] Page {page}:\n{smart_clean_text(ref_text)}\n")

        main_text += "\n\n" + "\n".join(refs)

    return main_text


def read_fb2(path: Path):
    content = decode_bytes_smart(path.read_bytes())
    soup = BeautifulSoup(content, "xml")
    return soup.get_text("\n", strip=True)


def read_fb2_zip(path: Path):
    with zipfile.ZipFile(path, "r") as z:
        fb2_files = [n for n in z.namelist() if n.lower().endswith(".fb2")]

        if not fb2_files:
            return ""

        with z.open(fb2_files[0]) as f:
            content = decode_bytes_smart(f.read())
            soup = BeautifulSoup(content, "xml")
            return soup.get_text("\n", strip=True)


def read_file(path: Path, ocr_langs: str):
    name = path.name.lower()
    suffix = path.suffix.lower()

    if name.endswith(".fb2.zip"):
        return read_fb2_zip(path)

    if suffix == ".txt":
        return read_txt(path)

    if suffix == ".md":
        return read_md(path)

    if suffix == ".rtf":
        return read_rtf(path)

    if suffix == ".pdf":
        return read_pdf_structural(path, ocr_langs)

    if suffix == ".docx":
        return read_docx(path)

    if suffix in [".xls", ".xlsx"]:
        return read_excel(path)

    if suffix == ".csv":
        return read_csv_file(path)

    if suffix == ".fb2":
        return read_fb2(path)

    if suffix == ".epub":
        return read_epub(path)

    if suffix in [".html", ".htm"]:
        return read_html(path)

    raise RuntimeError(f"Unsupported standalone format: {suffix}")


class FileMergerApp:
    def __init__(self, root):
        self.root = root
        self.lang = "en"
        self.t = LANG[self.lang]

        self.root.title(f"{APP_NAME} {APP_VERSION}")
        self.root.geometry("620x940")
        self.root.configure(bg="#0a0c10")

        if ICON_FILE.exists():
            try:
                self.root.iconbitmap(str(ICON_FILE))
            except Exception:
                pass

        self.files = []
        self.queue = queue.Queue()
        self.activity_step = 0
        self.is_working = False
        self.more_langs_visible = False
        self.ocr_vars = {}

        self.top_bar = tk.Frame(root, bg="#0a0c10")
        self.top_bar.pack(fill="x", padx=12, pady=(8, 0))

        self.lang_btn = tk.Button(
            self.top_bar,
            text="RU",
            command=self.toggle_language,
            width=6,
            height=1,
            bg="#15181c",
            fg="#00ff88",
            activebackground="#1f2933",
            activeforeground="#ffffff",
            bd=0,
            font=("Segoe UI", 9, "bold")
        )
        self.lang_btn.pack(side="right")

        self.title = tk.Label(
            root,
            text=f"{APP_NAME}\n{APP_VERSION}",
            font=("Segoe UI", 15, "bold"),
            bg="#0a0c10",
            fg="#00ff88",
            justify="center"
        )
        self.title.pack(pady=6)

        self.info = tk.Label(
            root,
            text=self.t["subtitle"],
            font=("Segoe UI", 9),
            bg="#0a0c10",
            fg="#8b949e"
        )
        self.info.pack()

        self.formats = tk.Label(
            root,
            text=self.t["formats"],
            font=("Segoe UI", 8),
            bg="#0a0c10",
            fg="#8b949e",
            justify="center"
        )
        self.formats.pack(pady=4)

        self.ocr_frame = tk.Frame(root, bg="#0a0c10")
        self.ocr_frame.pack(pady=4)

        self.ocr_title = tk.Label(
            self.ocr_frame,
            text=self.t["ocr_title"],
            font=("Segoe UI", 9, "bold"),
            bg="#0a0c10",
            fg="#00ff88"
        )
        self.ocr_title.pack()

        self.ocr_main_frame = tk.Frame(self.ocr_frame, bg="#0a0c10")
        self.ocr_main_frame.pack()

        installed_langs = get_installed_ocr_languages()
        installed_codes = {code for _, code in installed_langs}

        for label, code in OCR_LANG_OPTIONS_MAIN.items():
            if installed_codes and code not in installed_codes:
                continue

            var = tk.BooleanVar(value=True)
            self.ocr_vars[code] = var

            cb = tk.Checkbutton(
                self.ocr_main_frame,
                text=label,
                variable=var,
                bg="#0a0c10",
                fg="#e5e7eb",
                selectcolor="#15181c",
                activebackground="#0a0c10",
                activeforeground="#00ff88",
                font=("Segoe UI", 8)
            )
            cb.pack(side="left", padx=5)

        self.more_lang_btn = tk.Button(
            self.ocr_frame,
            text=self.t["more_langs"],
            command=self.toggle_more_ocr_langs,
            width=24,
            height=1,
            bg="#15181c",
            fg="#00ff88",
            bd=0,
            font=("Segoe UI", 8, "bold")
        )
        self.more_lang_btn.pack(pady=4)

        self.ocr_more_container = tk.Frame(self.ocr_frame, bg="#0a0c10")

        self.ocr_canvas = tk.Canvas(
            self.ocr_more_container,
            width=560,
            height=190,
            bg="#0a0c10",
            highlightthickness=0
        )

        self.ocr_scrollbar = tk.Scrollbar(
            self.ocr_more_container,
            orient="vertical",
            command=self.ocr_canvas.yview
        )

        self.ocr_scrollable_frame = tk.Frame(self.ocr_canvas, bg="#0a0c10")

        self.ocr_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.ocr_canvas.configure(
                scrollregion=self.ocr_canvas.bbox("all")
            )
        )

        self.ocr_canvas.create_window(
            (0, 0),
            window=self.ocr_scrollable_frame,
            anchor="nw"
        )

        self.ocr_canvas.configure(yscrollcommand=self.ocr_scrollbar.set)

        self.ocr_more_container.bind("<Enter>", self.bind_ocr_mousewheel)
        self.ocr_more_container.bind("<Leave>", self.unbind_ocr_mousewheel)
        self.ocr_canvas.bind("<Enter>", self.bind_ocr_mousewheel)
        self.ocr_canvas.bind("<Leave>", self.unbind_ocr_mousewheel)
        self.ocr_scrollable_frame.bind("<Enter>", self.bind_ocr_mousewheel)

        self.ocr_canvas.pack(side="left", fill="both", expand=True)
        self.ocr_scrollbar.pack(side="right", fill="y")

        more_langs = [
            (name, code)
            for name, code in installed_langs
            if code not in {"eng", "rus", "spa"}
        ]

        row = 0
        col = 0

        for name, code in more_langs:
            var = tk.BooleanVar(value=False)
            self.ocr_vars[code] = var

            cb = tk.Checkbutton(
                self.ocr_scrollable_frame,
                text=name,
                variable=var,
                bg="#0a0c10",
                fg="#e5e7eb",
                selectcolor="#15181c",
                activebackground="#0a0c10",
                activeforeground="#00ff88",
                font=("Segoe UI", 8)
            )
            cb.bind("<Enter>", self.bind_ocr_mousewheel)
            cb.grid(row=row, column=col, padx=6, pady=2, sticky="w")

            col += 1
            if col >= 3:
                col = 0
                row += 1

        self.select_btn = tk.Button(
            root,
            text=self.t["select"],
            command=self.select_files,
            width=24,
            height=2,
            bg="#15181c",
            fg="#00ff88",
            bd=0,
            font=("Segoe UI", 10, "bold")
        )
        self.select_btn.pack(pady=6)

        self.listbox = tk.Listbox(
            root,
            width=74,
            height=9,
            bg="#111827",
            fg="#e5e7eb",
            bd=0,
            font=("Consolas", 9)
        )
        self.listbox.pack(pady=6)

        self.convert_btn = tk.Button(
            root,
            text=self.t["convert"],
            command=self.start_conversion,
            width=24,
            height=2,
            bg="#00ff88",
            fg="#000000",
            bd=0,
            font=("Segoe UI", 10, "bold")
        )
        self.convert_btn.pack(pady=6)

        self.progress = ttk.Progressbar(
            root,
            orient="horizontal",
            length=480,
            mode="determinate"
        )
        self.progress.pack(pady=6)

        self.progress_label = tk.Label(
            root,
            text=f'{self.t["progress"]}: 0%',
            font=("Segoe UI", 9),
            bg="#0a0c10",
            fg="#8b949e"
        )
        self.progress_label.pack()

        self.status_label = tk.Label(
            root,
            text=self.t["idle"],
            font=("Segoe UI", 9, "bold"),
            bg="#0a0c10",
            fg="#00ff88"
        )
        self.status_label.pack(pady=2)

        self.log = tk.Text(
            root,
            height=9,
            width=74,
            bg="#020617",
            fg="#00ff88",
            bd=0,
            font=("Consolas", 9)
        )
        self.log.pack(pady=8)

        self.footer = tk.Label(
            root,
            text=APP_COPYRIGHT,
            font=("Segoe UI", 8),
            bg="#0a0c10",
            fg="#666666"
        )
        self.footer.pack(pady=4)

        self.log_message(self.t["ready"])
        self.root.after(100, self.process_queue)
        self.root.after(300, self.animate_activity)

    def bind_ocr_mousewheel(self, _event=None):
        self.root.bind_all("<MouseWheel>", self.on_ocr_mousewheel)
        self.root.bind_all("<Button-4>", self.on_ocr_mousewheel)
        self.root.bind_all("<Button-5>", self.on_ocr_mousewheel)

    def unbind_ocr_mousewheel(self, _event=None):
        self.root.unbind_all("<MouseWheel>")
        self.root.unbind_all("<Button-4>")
        self.root.unbind_all("<Button-5>")

    def on_ocr_mousewheel(self, event):
        if getattr(event, "num", None) == 4:
            direction = -1
        elif getattr(event, "num", None) == 5:
            direction = 1
        else:
            direction = -1 * int(event.delta / 120)

        self.ocr_canvas.yview_scroll(direction, "units")

    def get_selected_ocr_langs(self):
        selected = [
            code for code, var in self.ocr_vars.items()
            if var.get()
        ]
        return "+".join(selected)

    def toggle_more_ocr_langs(self):
        self.more_langs_visible = not self.more_langs_visible

        if self.more_langs_visible:
            self.ocr_more_container.pack(pady=4)
            self.more_lang_btn.config(text=self.t["hide_langs"])
        else:
            self.ocr_more_container.pack_forget()
            self.more_lang_btn.config(text=self.t["more_langs"])

    def toggle_language(self):
        if self.lang == "en":
            self.lang = "ru"
        elif self.lang == "ru":
            self.lang = "es"
        else:
            self.lang = "en"

        self.t = LANG[self.lang]
        self.apply_language()

    def apply_language(self):
        next_lang = "RU" if self.lang == "en" else "ES" if self.lang == "ru" else "EN"
        self.lang_btn.config(text=next_lang)

        self.info.config(text=self.t["subtitle"])
        self.formats.config(text=self.t["formats"])
        self.ocr_title.config(text=self.t["ocr_title"])
        self.more_lang_btn.config(text=self.t["hide_langs"] if self.more_langs_visible else self.t["more_langs"])
        self.select_btn.config(text=self.t["select"])
        self.convert_btn.config(text=self.t["convert"])

        if not self.is_working:
            self.status_label.config(text=self.t["idle"])
            self.progress_label.config(text=f'{self.t["progress"]}: 0%')

    def log_message(self, msg):
        self.log.insert(tk.END, msg + "\n")
        self.log.see(tk.END)

    def animate_activity(self):
        if self.is_working:
            dots = "." * ((self.activity_step % 3) + 1)
            self.status_label.config(text=f'{self.t["working"]}{dots}')
            self.activity_step += 1

        self.root.after(500, self.animate_activity)

    def reset_session(self):
        self.files = []
        self.listbox.delete(0, tk.END)
        self.log.delete("1.0", tk.END)
        self.progress.stop()
        self.progress.config(mode="determinate")
        self.progress["value"] = 0
        self.progress_label.config(text=f'{self.t["progress"]}: 0%')
        self.status_label.config(text=self.t["idle"])
        self.is_working = False
        self.log_message(self.t["new_session"])

    def select_files(self):
        patterns = " ".join(STANDALONE_FORMATS)

        selected = filedialog.askopenfilenames(
            title=self.t["select"],
            filetypes=[
                (self.t["supported_files"], patterns),
                (self.t["all_files"], "*.*"),
            ]
        )

        self.files = sorted([Path(x) for x in selected], key=natural_key)

        self.listbox.delete(0, tk.END)

        for i, file in enumerate(self.files, start=1):
            self.listbox.insert(tk.END, f"{i:03}. {file.name}")

        self.log_message(f'{self.t["selected"]}: {len(self.files)}')

    def start_conversion(self):
        if not self.files:
            messagebox.showwarning(self.t["no_files_title"], self.t["no_files_msg"])
            return

        ocr_langs = self.get_selected_ocr_langs()

        if not ocr_langs:
            messagebox.showwarning(self.t["no_ocr_title"], self.t["no_ocr_msg"])
            return

        out_file = filedialog.asksaveasfilename(
            title=self.t["save_title"],
            defaultextension=".txt",
            initialfile=make_output_filename(self.files),
            filetypes=[(self.t["text_files"], "*.txt")]
        )

        if not out_file:
            return

        self.select_btn.config(state="disabled")
        self.convert_btn.config(state="disabled")

        self.progress.config(mode="indeterminate")
        self.progress.start(12)
        self.progress_label.config(
            text=f'{self.t["progress"]}: 0% | {self.t["remaining"]}: {len(self.files)}'
        )
        self.status_label.config(text=f'{self.t["working"]}...')
        self.is_working = True

        self.log_message(f"OCR: {ocr_langs}")

        thread = threading.Thread(
            target=self.worker,
            args=(self.files.copy(), Path(out_file), ocr_langs),
            daemon=True
        )
        thread.start()

    def worker(self, files, out_file, ocr_langs):
        total = len(files)

        try:
            with open(out_file, "w", encoding="utf-8") as out:
                for i, file in enumerate(files, start=1):
                    self.queue.put(("log", f'{i:03}. {self.t["adding"]}: {file.name}'))

                    try:
                        text = read_file(file, ocr_langs)
                        text = smart_clean_text(text)

                        if text:
                            out.write(f"\n\n===== {i}. {file.name} =====\n\n")
                            out.write(text)
                            out.write("\n")
                        else:
                            self.queue.put(("log", f'{self.t["empty"]}: {file.name}'))

                    except Exception as e:
                        self.queue.put(("log", f'{self.t["skipped"]}: {file.name}: {e}'))

                    pct = int(i / total * 100)
                    remain = total - i
                    self.queue.put(("progress", pct, remain))

            self.queue.put(("done", str(out_file)))

        except Exception as e:
            self.queue.put(("error", str(e)))

    def process_queue(self):
        try:
            while True:
                msg = self.queue.get_nowait()

                if msg[0] == "log":
                    self.log_message(msg[1])

                elif msg[0] == "progress":
                    self.progress.stop()
                    self.progress.config(mode="determinate")
                    self.progress["value"] = msg[1]
                    self.progress_label.config(
                        text=f'{self.t["progress"]}: {msg[1]}% | {self.t["remaining"]}: {msg[2]}'
                    )

                    if msg[2] > 0:
                        self.progress.config(mode="indeterminate")
                        self.progress.start(12)

                elif msg[0] == "done":
                    out_file = msg[1]

                    self.is_working = False
                    self.progress.stop()
                    self.progress.config(mode="determinate")
                    self.progress["value"] = 100
                    self.progress_label.config(text=f'{self.t["progress"]}: 100%')
                    self.status_label.config(text=self.t["completed"])

                    self.select_btn.config(state="normal")
                    self.convert_btn.config(state="normal")

                    open_folder = messagebox.askyesno(
                        self.t["done"],
                        f'{self.t["file_converted"]}:\n{out_file}\n\n{self.t["open_folder"]}'
                    )

                    if open_folder:
                        subprocess.run(["explorer", "/select,", out_file])

                    again = messagebox.askyesno(
                        self.t["next"],
                        self.t["convert_more"]
                    )

                    if again:
                        self.reset_session()
                    else:
                        self.root.destroy()

                elif msg[0] == "error":
                    self.is_working = False
                    self.progress.stop()
                    self.progress.config(mode="determinate")
                    self.status_label.config(text=self.t["error"])

                    self.select_btn.config(state="normal")
                    self.convert_btn.config(state="normal")

                    messagebox.showerror(self.t["error"], msg[1])

        except queue.Empty:
            pass

        self.root.after(100, self.process_queue)


if __name__ == "__main__":
    root = tk.Tk()
    app = FileMergerApp(root)
    root.mainloop()
